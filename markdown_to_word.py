#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Document Converter

This script converts Markdown files to Word (.docx) format
for easy upload to Confluence.

Requirements:
    pip install python-docx markdown

Usage:
    python markdown_to_word.py input.md output.docx
    python markdown_to_word.py PROJECT_SUMMARY.md
"""

import sys
import os
import io
import re

# Don't wrap stdout here - it causes issues when imported as a module
# Only wrap if running as main script
if __name__ == '__main__' and sys.platform == 'win32':
    try:
        import io
        if hasattr(sys.stdout, 'buffer') and not isinstance(sys.stdout, io.TextIOWrapper):
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        if hasattr(sys.stderr, 'buffer') and not isinstance(sys.stderr, io.TextIOWrapper):
            sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except (AttributeError, TypeError, ValueError):
        pass

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed!")
    print("Please install it with: pip install python-docx")
    sys.exit(1)

try:
    import markdown
except ImportError:
    print("Error: markdown not installed!")
    print("Please install it with: pip install markdown")
    sys.exit(1)


def parse_markdown_to_docx(markdown_text, doc):
    """Parse Markdown text and add to Word document"""
    lines = markdown_text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Headers
        if stripped.startswith('# '):
            # h1
            text = stripped[2:].strip()
            para = doc.add_heading(text, level=1)
            i += 1
        elif stripped.startswith('## '):
            # h2
            text = stripped[3:].strip()
            para = doc.add_heading(text, level=2)
            i += 1
        elif stripped.startswith('### '):
            # h3
            text = stripped[4:].strip()
            para = doc.add_heading(text, level=3)
            i += 1
        elif stripped.startswith('#### '):
            # h4
            text = stripped[5:].strip()
            para = doc.add_heading(text, level=4)
            i += 1
        elif stripped.startswith('##### '):
            # h5
            text = stripped[6:].strip()
            para = doc.add_heading(text, level=5)
            i += 1
        elif stripped.startswith('###### '):
            # h6
            text = stripped[7:].strip()
            para = doc.add_heading(text, level=6)
            i += 1
        # Horizontal rule
        elif stripped == '---' or stripped == '***':
            para = doc.add_paragraph()
            para.add_run('_' * 80)
            i += 1
        # Empty line
        elif not stripped:
            doc.add_paragraph()
            i += 1
        # Code block
        elif stripped.startswith('```'):
            # Collect code block
            language = stripped[3:].strip() if len(stripped) > 3 else ''
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # Skip closing ```
            
            code_text = '\n'.join(code_lines)
            para = doc.add_paragraph(style='Intense Quote')
            run = para.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        # List item
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            # Remove bold/italic markers for now (Word will handle formatting)
            para = doc.add_paragraph(text, style='List Bullet')
            i += 1
        # Numbered list
        elif re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            para = doc.add_paragraph(text, style='List Number')
            i += 1
        # Table (basic support)
        elif '|' in line and stripped.startswith('|'):
            # Collect table
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                if not lines[i].strip().startswith('|---'):
                    table_lines.append(lines[i])
                i += 1
            
            if table_lines:
                # Parse table
                rows = []
                for table_line in table_lines:
                    cells = [cell.strip() for cell in table_line.split('|') if cell.strip()]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    # Create Word table
                    num_cols = len(rows[0])
                    num_rows = len(rows)
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if row_idx < num_rows and col_idx < num_cols:
                                table.rows[row_idx].cells[col_idx].text = cell_text
                                # Make header row bold
                                if row_idx == 0:
                                    for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
        # Regular paragraph
        else:
            # Process inline formatting
            para = doc.add_paragraph()
            add_formatted_text(para, line)
            i += 1


def add_formatted_text(para, text):
    """Add text with formatting (bold, italic, code)"""
    # Simple approach: parse basic markdown inline
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # Code
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # Italic (but not if it's part of **)
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            para.add_run(part)


def convert_markdown_to_word(input_file, output_file):
    """Convert Markdown file to Word document"""
    # Read Markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = os.path.splitext(os.path.basename(input_file))[0]
    
    # Parse and add content
    parse_markdown_to_docx(markdown_content, doc)
    
    # Save document
    doc.save(output_file)
    return True


def main():
    if len(sys.argv) < 2:
        print("Usage: python markdown_to_word.py <input.md> [output.docx]")
        print("\nExample:")
        print("  python markdown_to_word.py PROJECT_SUMMARY.md")
        print("  python markdown_to_word.py PROJECT_SUMMARY.md output.docx")
        print("\nRequirements:")
        print("  pip install python-docx markdown")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    if not os.path.exists(input_file):
        print(f"Error: File '{input_file}' not found.")
        sys.exit(1)
    
    # Determine output file
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        # Auto-generate output filename
        base_name = os.path.splitext(input_file)[0]
        output_file = f"{base_name}.docx"
    
    # Check if output file exists
    if os.path.exists(output_file):
        response = input(f"File '{output_file}' already exists. Overwrite? (y/n): ")
        if response.lower() != 'y':
            print("Conversion cancelled.")
            sys.exit(0)
    
    try:
        print(f"Converting '{input_file}' to '{output_file}'...")
        convert_markdown_to_word(input_file, output_file)
        print(f"[OK] Successfully converted to '{output_file}'")
        print(f"\nNext steps:")
        print(f"1. Open '{output_file}' in Word to review")
        print(f"2. In Confluence, click 'Create' → 'Import Word Document'")
        print(f"3. Or drag and drop the .docx file into Confluence")
        print(f"4. Confluence will convert it automatically")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()

